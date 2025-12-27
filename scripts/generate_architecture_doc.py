from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# Files to include snippets from
FILES = [
    'README.md',
    'conftest.py',
    'base/basepage.py',
    'pages/practice_page.py',
    'flows/alert_flow.py',
    'tests/test_conceptDemo.py',
]


def read_file_safe(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f'Could not read {path}: {e}'


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.bold = True


def main():
    doc = Document()
    doc.core_properties.title = 'Automation Framework Architecture & Flow'

    # Title
    title = doc.add_heading(level=0)
    r = title.add_run('Automation Framework: Architecture & Code Flow')
    r.bold = True
    r.font.size = Pt(18)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph('This document summarizes the current architecture and code flow of the automation framework present in this repository. It is intended to help new contributors understand responsibilities, interactions, and test execution flow.')

    # Project structure
    doc.add_heading('Project Structure', level=1)
    structure = doc.add_paragraph()
    structure.add_run('Key folders and files:').bold = True
    items = [
        ('base/', 'Contains BasePage with low-level Selenium helpers (click, type, wait, etc.)'),
        ('pages/', 'Page objects; represent pages and page-specific behaviors'),
        ('steps/', 'Higher-level flows or step helpers (composition with pages)'),
        ('tests/', 'Pytest tests expressing intent using fixtures and pages'),
        ('conftest.py', 'Driver and page fixtures (setup / teardown)'),
        ('requirements.txt', 'External Python dependencies (selenium, pytest, webdriver-manager)'),
    ]
    for name, desc in items:
        paragraph = doc.add_paragraph(style='List Bullet')
        paragraph.add_run(f'{name}').bold = True
        paragraph.add_run(f' — {desc}')

    # Responsibilities
    doc.add_heading('Responsibility by Layer', level=1)
    doc.add_paragraph('Tests ➜ Pages ➜ BasePage ➜ WebDriver')
    responsibilities = [
        ('Tests', 'Express business intent and assertions; use fixtures to get page instances'),
        ('Pages', 'Encapsulate locators and page-specific actions; may delegate to Steps helpers (composition)'),
        ('Steps', 'Hold higher-level flows and use a page instance (composition) to orchestrate multiple page actions'),
        ('BasePage', 'Provide low-level interactions and common waits for reliable operations'),
        ('conftest', 'Fixtures for driver lifecycle and page instantiation'),
    ]
    for name, desc in responsibilities:
        p = doc.add_paragraph()
        p.add_run(name + ': ').bold = True
        p.add_run(desc)

    # Sequence / Execution Flow
    doc.add_heading('Execution Flow (Typical Test)', level=1)
    steps = [
        'Pytest collects tests and fixtures',
        '`driver` fixture starts a browser and navigates to base URL',
        '`practice_page` fixture constructs a `PracticePage(driver)` instance',
        'Test calls page methods (e.g., `practice_page.click_show_button()`)',
        'Page methods use BasePage helpers (e.g., `clickit`, `typeit`) to interact with WebDriver',
        'For complex flows, Page delegates to `PracticePageSteps` (composition) which uses the page instance',
        'Alerts / dialogs and other browser interactions are handled via BasePage helpers (e.g., `switch_to_alert`)',
        'Assertions verify expected behavior, and fixtures teardown closes the browser',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    # Current implementation notes
    doc.add_heading('Current Implementation Notes', level=1)
    doc.add_paragraph('Recent changes: `PracticePage` now instantiates `PracticePageSteps(self)` and delegates higher-level flows to it. This uses composition rather than inheritance for steps/flows.')

    # Architecture diagram (generated PNG)
    doc.add_heading('Architecture Diagram', level=1)
    # ensure diagram exists (generated below) and insert it
    diagram_path = os.path.join(ROOT, 'Architecture_diagram.png')
    if not os.path.exists(diagram_path):
        create_architecture_diagram(diagram_path)
    # insert diagram sized to fit
    try:
        from docx.shared import Inches
        doc.add_picture(diagram_path, width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f'Could not insert diagram: {e}')

    # How to run tests
    doc.add_heading('How to Run Tests', level=1)
    doc.add_paragraph('1. Create a virtual environment (recommended)')
    doc.add_paragraph('2. Install dependencies: `pip install -r requirements.txt`')
    doc.add_paragraph('3. Run tests: `python -m pytest -q`')

    # Tips & Next steps
    doc.add_heading('Tips & Next Steps', level=1)
    tips = [
        'Consider adding a README for each folder describing its intent',
        'Add automatic doc generation or a diagram export for visualization',
        'Add unit tests for Steps classes by injecting fake page objects to validate orchestration logic',
        'Consider adding a simple architecture diagram (PNG) into the docs folder',
    ]
    for t in tips:
        doc.add_paragraph(t, style='List Bullet')

    # Save
    out_path = os.path.join(ROOT, 'Architecture.docx')
    doc.save(out_path)
    print(f'Architecture document written to: {out_path}')


# === Detailed interview-oriented document generator ===

def generate_from_scratch_doc():
    doc2 = Document()
    doc2.core_properties.title = 'How We Built The Automation Framework (From Scratch)'

    # Title
    t2 = doc2.add_heading(level=0)
    r2 = t2.add_run('How We Built The Automation Framework — Step by Step')
    r2.bold = True
    r2.font.size = Pt(18)
    t2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Short introduction
    doc2.add_heading('Purpose', level=1)
    doc2.add_paragraph('This document is crafted to help you explain, in interviews or rehearsals, where and how this automation framework was prepared from scratch. It emphasises key design decisions, step-by-step implementation tasks, and suggested talking points for interviews.')

    # From scratch steps
    doc2.add_heading('From Scratch — Implementation Steps', level=1)
    steps = [
        ('1) Project initialization', 'Create a Git repo and Python virtual environment. Initialize basic project layout (folders: base, pages, flows, tests, scripts). Add `requirements.txt` and `README.md`.'),
        ('2) Add test runner', 'Choose pytest, add `pytest` to requirements; create `tests/` and a sample test to validate discovery.'),
        ('3) Driver fixture', 'Create `conftest.py` with a `driver` fixture that installs and starts Chrome via `webdriver-manager`, reads `--env` option and navigates to base URL.'),
        ('4) BasePage', 'Implement `BasePage` with low-level helpers (clickit, typeit, waits, alert handlers) to centralize Selenium interactions and reduce duplication.'),
        ('5) Page Objects', 'Add page classes in `pages/` (e.g., `PracticePage`) that store locators and expose intent-driven methods using `BasePage` helpers.'),
        ('6) Flows/Steps', 'Create `flows/` to hold higher-level orchestration (e.g., `AlertFlow`, `AutoSuggestionsFlow`) which compose page methods for complex scenarios.'),
        ('7) Tests and fixtures', 'Write tests that use page fixtures (avoid globals) and focus on business intent and clear assertions.'),
        ('8) Debugging & Stability', 'Add explicit waits, handle visibility vs presence correctly, and identify/pinpoint flakiness with logs and screenshots (if added).'),
        ('9) Documentation', 'Document the architecture in `README.md` and generate an exportable doc (`Framework_From_Scratch.docx`) for interviews.'),
    ]
    for title_text, desc in steps:
        p = doc2.add_paragraph(style='List Number')
        p.add_run(title_text + ': ').bold = True
        p.add_run(desc)

    # Where (file map)
    doc2.add_heading('Where to find things (File map)', level=1)
    mappings = [
        ('`conftest.py`', 'Driver and environment fixtures, `--env` test option'),
        ('`base/basepage.py`', 'Low-level Selenium utils and waits — use this to explain how you prevented flakiness'),
        ('`pages/*.py`', 'Page Objects that express page behavior and locators'),
        ('`flows/*.py`', 'Domain flows / composition to orchestrate multiple page actions'),
        ('`tests/*.py`', 'Intent-focused tests; examples of using flows and page fixtures'),
        ('`scripts/generate_architecture_doc.py`', 'Generates Word doc with architecture and snippets (this file)')
    ]
    for f, d in mappings:
        pr = doc2.add_paragraph()
        pr.add_run(f + ': ').bold = True
        pr.add_run(d)

    # How we implemented (detailed notes)
    doc2.add_heading('How we implemented — Notes & Rationale', level=1)
    notes = [
        ('Fixture design', 'Used `yield` fixtures to ensure reliable teardown. Avoided globals and made pages ephemeral per test to prevent state leakage.'),
        ('BasePage responsibilities', 'Holds waits and safe interactions; prefer explicit waits (`WebDriverWait`) and helper wrappers to avoid flakiness.'),
        ('Page Objects design', 'Pages keep locators and intent-level methods; tests should NOT access locators directly.'),
        ('Flows (composition)', 'Flows hold orchestration logic and compose page objects. This keeps pages small and focused on single page behavior.'),
        ('Debugging mistakes to mention', 'Missed parentheses on methods, recursive self-calls, and not waiting for visibility are realistic sources of early bugs.'),
    ]
    for t, d in notes:
        pr = doc2.add_paragraph()
        pr.add_run(t + ': ').bold = True
        pr.add_run(d)

    # Interview talking points
    doc2.add_heading('Interview Talking Points (What to say)', level=1)
    bullets = [
        'Explain why separation of concerns (Tests → Pages → BasePage) increases maintainability.',
        'Discuss why explicit waits (WebDriverWait) are used and the difference between presence vs visibility.',
        'Talk about fixture lifecycles and why `yield` is preferred for setup/teardown control.',
        'Describe the decision to use composition (Flows) instead of large page classes or inheritance for complex flows.',
        'Emphasize how small, intent-driven tests improve readability and reduce brittleness.'
    ]
    for b in bullets:
        doc2.add_paragraph(b, style='List Bullet')

    # Common interview Q&A (short sample answers)
    doc2.add_heading('Sample Q&A (short answers you can rehearse)', level=1)
    qas = [
        ("Why a BasePage?", "Centralizes interaction methods and waits so Pages/tests remain concise and stable."),
        ("How do you avoid flaky tests?", "Use explicit waits, minimize sleep usage, keep independent test state, and use page flows to reduce UI timing fragility."),
        ("Why use fixtures?", "Fixtures manage resource lifecycle and configuration, enabling reusable setup (driver, login) while ensuring proper teardown."),
        ("When to use a Flow?", "When a user journey spans multiple pages or complex orchestration — keep that logic outside a single Page object."),
    ]
    for q, a in qas:
        p = doc2.add_paragraph()
        p.add_run(q + ' ').bold = True
        p.add_run(a)

    # How to run (practical commands)
    doc2.add_heading('How to run — quick commands', level=1)
    run_cmds = [
        ('Create venv', 'python -m venv .venv'),
        ('Activate venv (Windows)', '.venv\\Scripts\\activate'),
        ('Install deps', 'pip install -r requirements.txt'),
        ('Run tests', 'python -m pytest -q'),
        ('Run a single test', 'python -m pytest tests/test_conceptDemo.py::test_check_alert_message_content -q'),
        ('Run with env', 'python -m pytest -q --env=uat')
    ]
    for title_text, cmd in run_cmds:
        p = doc2.add_paragraph(style='List Number')
        p.add_run(title_text + ': ').bold = True
        p.add_run(cmd)

    # Appendix: include useful snippets from source files
    doc2.add_heading('Appendix — Useful Code Snippets', level=1)
    for f in FILES:
        content = read_file_safe(os.path.join(ROOT, f))
        doc2.add_heading(f, level=2)
        para = doc2.add_paragraph()
        run = para.add_run(content)
        try:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        except Exception:
            pass

    # Save detailed doc
    out_path = os.path.join(ROOT, 'Framework_From_Scratch.docx')
    doc2.save(out_path)

    print(f'Framework-from-scratch document written to: {out_path}')


# call the detailed doc generator when executed
generate_from_scratch_doc()


# --- Diagram generation using Pillow ---
from PIL import Image, ImageDraw, ImageFont

def create_architecture_diagram(out_path):
    """Create a simple OO-style architecture diagram and save as PNG."""
    W, H = 1200, 600
    img = Image.new('RGB', (W, H), color='white')
    d = ImageDraw.Draw(img)
    # fonts
    try:
        font_bold = ImageFont.truetype('arialbd.ttf', 18)
        font = ImageFont.truetype('arial.ttf', 14)
    except Exception:
        font_bold = ImageFont.load_default()
        font = ImageFont.load_default()

    # box positions
    boxes = {
        'Tests': (480, 20, 720, 80),
        'Pages': (80, 140, 360, 220),
        'Steps': (440, 140, 760, 220),
        'BasePage': (80, 300, 360, 380),
        'WebDriver': (480, 300, 760, 380),
    }

    # draw boxes
    for label, (x1, y1, x2, y2) in boxes.items():
        d.rectangle([x1, y1, x2, y2], outline='black', width=2, fill='#f3f4f6')
        bbox = d.textbbox((0, 0), label, font=font_bold)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        d.text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), label, fill='black', font=font_bold)

    # arrows (Tests -> Pages, Tests -> Steps)
    def arrow(a, b):
        ax = (a[0] + a[2]) / 2
        ay = a[3]
        bx = (b[0] + b[2]) / 2
        by = b[1]
        d.line((ax, ay, bx, by), fill='black', width=2)
        # arrowhead
        d.polygon([(bx-6, by-10), (bx+6, by-10), (bx, by)], fill='black')

    arrow(boxes['Tests'], boxes['Pages'])
    arrow(boxes['Tests'], boxes['Steps'])
    arrow(boxes['Pages'], boxes['BasePage'])
    arrow(boxes['Steps'], boxes['BasePage'])
    arrow(boxes['BasePage'], boxes['WebDriver'])

    # labels for responsibilities
    d.text((420, 90), 'Test execution → orchestrates pages and asserts behavior', fill='black', font=font)
    d.text((420, 110), 'Pages delegate to Steps for complex flows; BasePage holds low-level interactions', fill='black', font=font)

    img.save(out_path)



if __name__ == '__main__':
    main()
